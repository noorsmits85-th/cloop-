import os
import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
                <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

doc = docx.Document()

# 1. Căn lề chuẩn: Trên 2cm, Dưới 2cm, Trái 3cm, Phải 2cm
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

# 2. Cấu hình Font mặc định: Times New Roman, size 13pt
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
style.paragraph_format.line_spacing = 1.25
style.paragraph_format.space_after = Pt(4)

def add_doc_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0A, 0x25, 0x17)

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x4A, 0x34)

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x18, 0x3A, 0x2D)

def add_body_p(text, bold_prefix=None, italic_suffix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Times New Roman'
        r_bold.font.size = Pt(13)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    r_text = p.add_run(text)
    r_text.font.name = 'Times New Roman'
    r_text.font.size = Pt(13)
    if italic_suffix:
        r_it = p.add_run(italic_suffix)
        r_it.font.name = 'Times New Roman'
        r_it.font.size = Pt(13)
        r_it.font.italic = True
    return p

def add_bullet_item(bold_label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    r_bold = p.add_run(bold_label)
    r_bold.font.name = 'Times New Roman'
    r_bold.font.size = Pt(13)
    r_bold.font.bold = True
    r_text = p.add_run(text)
    r_text.font.name = 'Times New Roman'
    r_text.font.size = Pt(13)

# ==================== NỘI DUNG TÀI LIỆU ====================

add_doc_title("BÁO CÁO KIẾN TRÚC HỆ THỐNG VẬN HÀNH & KINH TẾ HỌC NỀN TẢNG CLOOP")

# --- I. KIẾN TRÚC HỆ THỐNG THỜI GIAN THỰC ---
add_heading_1("I. KIẾN TRÚC HỆ THỐNG THỜI GIAN THỰC")
add_body_p("Nền tảng vận hành trên cơ sở hạ tầng API trực tiếp, đảm bảo tính tự động hóa và toàn vẹn dữ liệu:")
add_bullet_item("Logistics Động (GHN Gateway): ", "Tích hợp máy chủ Giao Hàng Nhanh (online-gateway.ghn.vn). Hệ thống tự động truy xuất bản đồ hành chính, đo khoảng cách thực tế giữa Tủ đồ (Người gửi) và Khách thuê (Người nhận) để định tuyến và áp biểu phí thời gian thực.")
add_bullet_item("Thanh toán & Két Escrow (PayOS): ", "Tích hợp cổng thanh toán trực tuyến. Mỗi đơn hàng sinh một mã VietQR động định danh chính xác số tiền lẻ đến từng đồng, khóa an toàn dòng tiền trong Két Escrow cho đến khi đơn hàng hoàn tất.")
add_bullet_item("Quản trị Dữ liệu (Supabase / PostgreSQL): ", "Sử dụng cơ chế Khóa bi quan (Pessimistic Locking) ở cấp độ hàng (row-level) để loại bỏ 100% rủi ro đụng lịch thuê.")
add_bullet_item("Tối ưu Hóa Hình ảnh (Cloudinary CDN): ", "Tự động nén và phân phối dữ liệu trang phục qua mạng lưới CDN, duy trì tốc độ tải trang dưới 1 giây.")

# --- II. THUẬT TOÁN LOGISTICS SAN SẺ 50/50 & ĐỆM RỦI RO BLOCK 5K ---
add_heading_1("II. THUẬT TOÁN LOGISTICS SAN SẺ 50/50 & ĐỆM RỦI RO BLOCK 5K")
add_body_p("Nhằm tối ưu hóa tỷ lệ chuyển đổi (CRO), hạ thấp rào cản chi phí cho khách hàng và bảo toàn nguồn vốn, CLOOP áp dụng mô hình San sẻ trách nhiệm vận chuyển (Co-sharing Logistics) kết hợp Thuật toán Đệm Block 5K (Safety Buffer):")

add_heading_2("1. Thuật toán Làm tròn Block 5K")
add_body_p("Mọi mức cước phát sinh từ API GHN đều được chuẩn hóa qua hàm trần (Ceil) lên mốc 5.000đ gần nhất:")
p_formula = doc.add_paragraph()
p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_formula.paragraph_format.space_before = Pt(3)
p_formula.paragraph_format.space_after = Pt(6)
r_f = p_formula.add_run("Cước niêm yết = ⌈ Cước thực tế / 5.000 ⌉ × 5.000")
r_f.font.name = 'Times New Roman'
r_f.font.size = Pt(13)
r_f.font.bold = True
r_f.font.color.rgb = RGBColor(0x0F, 0x4A, 0x34)

# Bảng Block 5K
table_block5k = doc.add_table(rows=5, cols=4)
table_block5k.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table_block5k)

headers_b5k = ["Tuyến vận chuyển", "Cước thực tế GHN", "Cước niêm yết Block 5K", "Khoản đệm & Ý nghĩa vận hành"]
data_b5k = [
    ["Nội thành tiêu chuẩn", "18.000₫", "20.000₫", "+2.000₫ (Bù hao hụt tuyến ngắn)"],
    ["Liên tỉnh nhẹ (<500g)", "21.000₫", "25.000₫", "+4.000₫ (Bù phát sinh giao lại lần 2)"],
    ["Huyện / Vùng xa", "22.500₫", "25.000₫", "+2.500₫ (Dự phòng phụ phí vùng sâu)"],
    ["Đầm tiệc cồng kềnh", "36.000₫", "40.000₫", "+4.000₫ (Bù sai số quy đổi thể tích D×R×C/5000)"]
]

hdr_cells = table_block5k.rows[0].cells
for i, name in enumerate(headers_b5k):
    hdr_cells[i].text = name
    set_cell_background(hdr_cells[i], "E6F4EA")
    set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0A, 0x25, 0x17)

for r_idx, row_data in enumerate(data_b5k):
    row_cells = table_block5k.rows[r_idx + 1].cells
    bg = "FFFFFF" if r_idx % 2 == 0 else "F9FBF9"
    for c_idx, val in enumerate(row_data):
        row_cells[c_idx].text = val
        set_cell_background(row_cells[c_idx], bg)
        set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
        p = row_cells[c_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [1, 2] else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

add_heading_2("2. Nguyên lý San sẻ 50/50")
add_bullet_item("Khách thuê (Chiều đi): ", "Thanh toán cước 1 chiều đã làm tròn (25.000đ) trực tiếp lúc quét mã VietQR.")
add_bullet_item("Chủ tủ (Chiều về): ", "Chịu cước 1 chiều đã làm tròn (25.000đ) như một khoản chi phí thu hồi và bảo toàn tài sản. Khoản này được hệ thống tự động cấn trừ (Auto-deduct) từ doanh thu trước khi giải ngân Payout.")
add_bullet_item("Quỹ Phòng Vệ Vận Chuyển: ", "Mỗi đơn hàng khứ hồi tạo ra khoản thặng dư +8.000đ (4.000đ × 2) được ghi nhận tự động vào Sổ cái (SHIPPING_RETAINED) để bảo chứng rủi ro giao lại hoặc sai lệch kích thước thùng carton.")

# --- III. CA KIỂM THỬ THỰC TẾ & MA TRẬN PHÂN BỔ DÒNG TIỀN ---
add_heading_1("III. CA KIỂM THỬ THỰC TẾ & MA TRẬN PHÂN BỔ DÒNG TIỀN")
add_body_p("Bối cảnh giao dịch mẫu trên môi trường Production:")
add_bullet_item("Chủ tủ: ", "Định vị tại Quận Cầu Giấy, Hà Nội.")
add_bullet_item("Khách thuê: ", "Định vị tại Huyện Hưng Nguyên, Nghệ An.")
add_bullet_item("Sản phẩm: ", "Váy dạ hội thiết kế (Gói thuê 3 ngày: 350.000đ, Tiền cọc Escrow: 1.000.000đ).")
add_bullet_item("Cước GHN thực tế: ", "21.000đ/chiều (Cước niêm yết áp dụng Block 5K: 25.000đ/chiều).")

add_body_p("Bảng Phân Bổ Dòng Tiền 4 Bên Qua Từng Giai Đoạn:")

table_flow = doc.add_table(rows=6, cols=5)
table_flow.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table_flow)

headers_flow = ["Giai đoạn luân chuyển", "Khách thuê", "Két Escrow CLOOP", "Chủ tủ đồ", "Hãng GHN"]
data_flow = [
    [
        "1. Quét mã VietQR Checkout",
        "-1.375.000₫",
        "+1.375.000₫\n(Thuê 350k + Cọc 1.000k + Ship đi 25k)",
        "0₫",
        "0₫"
    ],
    [
        "2. Bưu tá giao đồ (Lượt đi)",
        "0₫",
        "-21.000₫",
        "0₫",
        "+21.000₫"
    ],
    [
        "3. Thu hồi đồ (Lượt về Pre-paid)",
        "0₫",
        "-21.000₫",
        "0₫",
        "+21.000₫"
    ],
    [
        "4. Nghiệm thu & Payout Tiêu chuẩn",
        "+1.000.000₫\n(Hoàn 100% cọc)",
        "-1.000.000₫ (Nhả cọc)\n-283.000₫ (Payout)\n+42.000₫ (Phí sàn 12%)\n+8.000₫ (Quỹ phòng vệ)",
        "+283.000₫\n(350k - 42k sàn - 25k ship)",
        "0₫"
    ],
    [
        "4b. Nghiệm thu & Payout Founding 100",
        "+1.000.000₫\n(Hoàn 100% cọc)",
        "-1.000.000₫ (Nhả cọc)\n-325.000₫ (Payout)\n+0₫ (Phí sàn 0%)\n+8.000₫ (Quỹ phòng vệ)",
        "+325.000₫\n(350k - 0đ sàn - 25k ship)",
        "0₫"
    ]
]

hdr_flow = table_flow.rows[0].cells
for i, name in enumerate(headers_flow):
    hdr_flow[i].text = name
    set_cell_background(hdr_flow[i], "E6F4EA")
    set_cell_margins(hdr_flow[i], top=120, bottom=120, left=120, right=120)
    p = hdr_flow[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0A, 0x25, 0x17)

for r_idx, row_data in enumerate(data_flow):
    row_cells = table_flow.rows[r_idx + 1].cells
    bg = "FFFFFF" if r_idx % 2 == 0 else "F9FBF9"
    for c_idx, val in enumerate(row_data):
        row_cells[c_idx].text = val
        set_cell_background(row_cells[c_idx], bg)
        set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=120, right=120)
        p = row_cells[c_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in [0, 2, 3] else WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

p_conc = doc.add_paragraph()
p_conc.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_conc.paragraph_format.space_before = Pt(6)
p_conc.paragraph_format.space_after = Pt(8)
r_conc = p_conc.add_run("=> Tong dong tien Ket Escrow = Can bang so cai kep tuyet doi (Zero-Leakage)")
r_conc.font.name = 'Times New Roman'
r_conc.font.size = Pt(12.5)
r_conc.font.bold = True
r_conc.font.color.rgb = RGBColor(0x0F, 0x4A, 0x34)

# --- IV. CƠ CHẾ XỬ LÝ 3 TRƯỜNG HỢP BIÊN (EDGE CASES) ---
add_heading_1("IV. CƠ CHẾ XỬ LÝ 3 TRƯỜNG HỢP BIÊN (EDGE CASES)")

add_bullet_item("Trường hợp 1: Khách hủy đơn trước khi bàn giao cho GHN: ", "Hệ thống kích hoạt lệnh hoàn trả 100% toàn bộ số tiền (Tiền thuê + Tiền cọc + 25.000đ cước ship) về tài khoản của khách qua cổng ngân hàng.")
add_bullet_item("Trường hợp 2: Khách hủy đơn khi shipper GHN đã tiếp nhận hàng: ", "Khoản cước 25.000đ chiều đi bị giữ lại để thanh toán chi phí vận hành cho bưu chính; hệ thống chỉ hoàn trả lại tiền thuê và tiền cọc cho khách.")
add_bullet_item("Trường hợp 3: Phát sinh khiếu nại / Hư hỏng trang phục (Dispute): ", "Két Escrow đóng băng ngay lập tức khoản cọc 1.000.000đ. Phí vận chuyển 2 chiều thực tế (42.000đ) vẫn được tất toán sòng phẳng cho GHN. Hội đồng Trọng tài tiếp nhận biên bản video mở hộp trên /admin/disputes để ra phán quyết trích cọc bồi thường cho chủ tủ.")

# --- V. QUY TRÌNH VẬN HÀNH 4 GIAI ĐOẠN CỦA CHỦ TỦ (/my-closet) ---
add_heading_1("V. QUY TRÌNH VẬN HÀNH 4 GIAI ĐOẠN CỦA CHỦ TỦ (/my-closet)")
add_bullet_item("Giai đoạn 1: Tiếp nhận đơn & Đóng gói: ", "Chỉ khi nhận được Webhook từ PayOS báo trạng thái đã khóa 100% tiền cọc và tiền thuê, chủ tủ mới nhận thông báo đóng gói. Chủ tủ bấm 'Đã đóng gói & Gọi shipper' -> Hệ thống tự sinh mã vận đơn ẩn danh (Masked Shipping Label) bảo mật thông tin cá nhân.")
add_bullet_item("Giai đoạn 2: Bàn giao & Kích hoạt chu trình thuê: ", "Shipper GHN đến lấy hàng tận nhà chủ tủ. Đồng hồ thời gian thuê (1 ngày, 3 ngày, 7 ngày) chỉ chính thức đếm ngược khi trạng thái đơn hàng trên GHN chuyển sang Đã giao hàng thành công (Delivered).")
add_bullet_item("Giai đoạn 3: Thu hồi đồ tự động 1-Click: ", "Hết hạn thuê, hệ thống tự động kích hoạt mã vận đơn chiều về trả trước (Pre-paid). Khách bàn giao đồ cho bưu tá mà không phải trả tiền mặt. Chủ tủ theo dõi tiến trình bưu phẩm hoàn hồi về tủ đồ.")
add_bullet_item("Giai đoạn 4: Nghiệm thu & Nhận Payout tự động: ", "Sau khi nhận lại bưu phẩm, chủ tủ có 24 giờ để kiểm tra. Bấm 'Xác nhận nguyên vẹn' (hoặc hết 24h không có khiếu nại) -> Két Escrow nhả cọc 100% cho khách và tự động chuyển tiền Payout đã cấn trừ phí ship về tài khoản ngân hàng của chủ tủ.")

# --- VI. BÀI TOÁN KINH TẾ HỌC & CHIẾN DỊCH 'FOUNDING 100' ---
add_heading_1("VI. BÀI TOÁN KINH TẾ HỌC & CHIẾN DỊCH 'FOUNDING 100'")

add_heading_2("1. Bảng So Sánh Đặc Quyền Thành Viên")

table_f100 = doc.add_table(rows=5, cols=3)
table_f100.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table_f100)

headers_f100 = ["Tiêu chí", "Thành viên tiêu chuẩn", "Thành viên Founding 100"]
data_f100 = [
    ["Phí dịch vụ sàn (Take-rate)", "12% - 15% / đơn", "Miễn phí 0% (trong 3 - 6 tháng)"],
    ["Doanh thu thực nhận (Payout)", "Tiền thuê - 12% - 25k ship về", "Tiền thuê - 0% - 25k ship về"],
    ["Huy hiệu hồ sơ", "Thành viên mới", "Huy hiệu độc quyền Founding Closet ⭐"],
    ["Thuật toán hiển thị", "Xếp hạng tự nhiên", "Đẩy thẳng lên Top Đề Xuất Trang Chủ"]
]

hdr_f100 = table_f100.rows[0].cells
for i, name in enumerate(headers_f100):
    hdr_f100[i].text = name
    set_cell_background(hdr_f100[i], "E6F4EA")
    set_cell_margins(hdr_f100[i], top=120, bottom=120, left=150, right=150)
    p = hdr_f100[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0A, 0x25, 0x17)

for r_idx, row_data in enumerate(data_f100):
    row_cells = table_f100.rows[r_idx + 1].cells
    bg = "FFFFFF" if r_idx % 2 == 0 else "F9FBF9"
    for c_idx, val in enumerate(row_data):
        row_cells[c_idx].text = val
        set_cell_background(row_cells[c_idx], bg)
        set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
        p = row_cells[c_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

add_heading_2("2. Lộ Trình 3 Giai Đoạn Chuyển Đổi Sinh Lời")
add_bullet_item("Giai đoạn 1 (Thâm nhập & Tích lũy Nguồn cung): ", "Áp dụng 0% phí sàn để thu hút 100 chủ tủ chất lượng cao đầu tiên, thiết lập 100 giao dịch tuần hoàn chuẩn mực.")
add_bullet_item("Giai đoạn 2 (Tạo lập thói quen kinh doanh): ", "Chủ tủ hình thành thói quen tạo dòng tiền thụ động từ tủ đồ nhàn rỗi, gắn bó mật thiết với hạ tầng vận hành tự động của CLOOP.")
add_bullet_item("Giai đoạn 3 (Thu phí & Tạo lập lợi nhuận bền vững): ", "Áp dụng mức phí sàn 12%. Chủ tủ sẵn sàng chi trả vì nền tảng đã thay họ gánh vác toàn bộ rủi ro tài chính, điều phối vận chuyển và liên tục cung cấp tệp khách hàng mới.")

# Save files
project_path = r"c:\Users\Yoga gen 3\Downloads\CLOOP_Techfest\cloop-app\BAO_CAO_KIEN_TRUC_VAN_HANH_CLOOP.docx"
downloads_path = r"c:\Users\Yoga gen 3\Downloads\BAO_CAO_KIEN_TRUC_VAN_HANH_CLOOP.docx"

doc.save(project_path)
doc.save(downloads_path)

print(f"SUCCESS: File saved at {project_path} and {downloads_path}")
